from __future__ import annotations

import json
import shutil
import time
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from fastapi.testclient import TestClient
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from whatsapp_ai_agent.config import Settings, get_settings
from whatsapp_ai_agent.db.models import Base, Membership, Organization, User
from whatsapp_ai_agent.db.session import get_db_session
from whatsapp_ai_agent.documents.schemas import DocumentTableUpdateRequest
from whatsapp_ai_agent.llm.schemas import ChatParseResult
from whatsapp_ai_agent.main import create_app
from whatsapp_ai_agent.workflows import chat_processing

XLSX_HEADERS = [
    "Work Date",
    "Time In",
    "Time Out",
    "Task / Work Done",
    "Technician(s)",
    "Work Area",
    "Completion",
    "Remarks",
]

XLSX_SCENARIOS = [
    {
        "name": "xlsx_append_earthing_test",
        "sid": "SMCOMPLEX001",
        "body": (
            "Update the Styled Field Activity Register Excel. On 2026-07-09 from "
            "09:10 to 12:40, activity was earthing continuity test at Admin Block. "
            "People participated: Samuel and Grace. Status done. Notes: readings "
            "within limit."
        ),
        "expected": {
            "Work Date": "2026-07-09",
            "Time In": "09:10",
            "Time Out": "12:40",
            "Task / Work Done": "earthing continuity test",
            "Technician(s)": "Samuel and Grace",
            "Work Area": "Admin Block",
            "Completion": "done",
            "Remarks": "readings within limit",
        },
    },
    {
        "name": "xlsx_update_existing_complex_row",
        "sid": "SMCOMPLEX002",
        "body": (
            "Update existing row in Styled Field Activity Register Excel for "
            "2026-06-25, 13:00 to 15:30. Activity: cable ladder alignment. "
            "People: Tunde and Samuel. Site: Workshop. Status done. Notes: all "
            "brackets fitted."
        ),
        "expected": {
            "Work Date": "2026-06-25",
            "Time In": "13:00",
            "Time Out": "15:30",
            "Task / Work Done": "cable ladder alignment",
            "Technician(s)": "Tunde and Samuel",
            "Work Area": "Workshop",
            "Completion": "done",
            "Remarks": "all brackets fitted",
        },
    },
    {
        "name": "xlsx_append_non_consecutive_later_date",
        "sid": "SMCOMPLEX003",
        "body": (
            "Append to Styled Field Activity Register Excel: on 2026-07-18 from "
            "14:05 to 17:20, activity was UPS bypass test at Server Room. People "
            "participated: Musa, Samuel, and Ada. Status needs_review. Notes: "
            "voltage dip observed during transfer."
        ),
        "expected": {
            "Work Date": "2026-07-18",
            "Time In": "14:05",
            "Time Out": "17:20",
            "Task / Work Done": "UPS bypass test",
            "Technician(s)": "Musa, Samuel, and Ada",
            "Work Area": "Server Room",
            "Completion": "needs_review",
            "Remarks": "voltage dip observed during transfer",
        },
    },
]

DOCX_SCENARIOS = [
    {
        "name": "docx_update_existing_title_header_table",
        "sid": "SMCOMPLEX004",
        "body": (
            "Update the Transformer Closeout Word table. Asset TX-01, work date "
            "2026-06-19. People: Samuel and Kemi. Action taken: oil level checked "
            "and gasket inspected. Result closed. Remarks: top-up completed."
        ),
        "expected": [
            "TX-01",
            "2026-06-19",
            "Samuel and Kemi",
            "oil level checked and gasket inspected",
            "closed",
            "top-up completed",
        ],
    },
    {
        "name": "docx_append_new_asset_row",
        "sid": "SMCOMPLEX005",
        "body": (
            "Append to the Transformer Closeout Word table: Asset TX-02, work date "
            "2026-07-08. People: Grace. Action taken: insulation resistance test. "
            "Result open. Remarks: awaiting supervisor review."
        ),
        "expected": [
            "TX-02",
            "2026-07-08",
            "Grace",
            "insulation resistance test",
            "open",
            "awaiting supervisor review",
        ],
    },
]


def create_complex_workbook(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Field Register"
    notes = workbook.create_sheet("Instructions")
    notes["A1"] = "This extra sheet should not be selected for document updates."

    sheet.merge_cells("A1:H1")
    sheet["A1"] = "ABUAD Engineering Field Activity Register"
    sheet["A1"].font = Font(bold=True, color="FFFFFF", size=14)
    sheet["A1"].fill = PatternFill("solid", fgColor="1F4E78")
    sheet["A1"].alignment = Alignment(horizontal="center")
    sheet["A2"] = "Styled workbook with title rows, colors, a table, and aliases."
    sheet["A2"].font = Font(italic=True, color="666666")
    sheet.append([])
    sheet.append(XLSX_HEADERS)
    seed_rows = [
        [
            "2026-06-20",
            "08:00",
            "10:00",
            "Panel inspection",
            "Samuel, Ada",
            "LV Room",
            "done",
            "No abnormal heating",
        ],
        [
            "2026-06-25",
            "13:00",
            "15:30",
            "Cable ladder alignment",
            "Tunde",
            "Workshop",
            "in_progress",
            "Brackets pending",
        ],
    ]
    thin = Side(style="thin", color="A6A6A6")
    for row in seed_rows:
        sheet.append(row)
    for cell in sheet[4]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="305496")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=thin)
    for row in sheet.iter_rows(min_row=5, max_row=6, min_col=1, max_col=8):
        for cell in row:
            cell.fill = PatternFill("solid", fgColor="D9EAF7")
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    table = Table(displayName="FieldRegisterTable", ref="A4:H6")
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    sheet.add_table(table)
    sheet.freeze_panes = "A5"
    sheet.auto_filter.ref = "A4:H6"
    for column, width in zip("ABCDEFGH", [14, 12, 12, 28, 22, 18, 14, 30], strict=False):
        sheet.column_dimensions[column].width = width
    workbook.save(path)


def create_complex_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Maintenance Closeout Register", level=1)
    document.add_paragraph("The table has a merged title row before the real headers.")
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    title_cells = table.rows[0].cells
    title_cells[0].text = "Transformer Maintenance Table"
    for cell in title_cells[1:]:
        title_cells[0].merge(cell)
    title_run = title_cells[0].paragraphs[0].runs[0]
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    headers = ["Asset Tag", "Work Date", "Person(s)", "Action Taken", "Result", "Remarks"]
    header_cells = table.add_row().cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        header_cells[index].paragraphs[0].runs[0].font.bold = True
    first = table.add_row().cells
    for index, value in enumerate(
        ["TX-01", "2026-06-19", "Samuel", "Oil level checked", "Open", "Top-up pending"]
    ):
        first[index].text = value
    document.save(path)


def norm(value: object) -> str:
    return " ".join(str(value or "").strip().lower().split())


def contains_expected(actual: object, expected: object) -> bool:
    actual_norm = norm(actual)
    expected_norm = norm(expected)
    if expected_norm and not actual_norm:
        return False
    return expected_norm in actual_norm or actual_norm in expected_norm


def read_xlsx_rows(path: Path) -> tuple[list[str], list[dict[str, object]], str]:
    workbook = load_workbook(path)
    sheet = workbook["Field Register"]
    headers = [cell.value for cell in sheet[4]]
    rows = []
    for values in sheet.iter_rows(min_row=5, max_col=8, values_only=True):
        if any(value is not None for value in values):
            rows.append(
                {str(header): value for header, value in zip(headers, values, strict=False)}
            )
    return [str(header) for header in headers], rows, sheet.tables["FieldRegisterTable"].ref


def find_xlsx_row(rows: list[dict[str, object]], date: str, start: str) -> dict[str, object] | None:
    for row in rows:
        date_matches = contains_expected(row.get("Work Date"), date)
        start_matches = contains_expected(row.get("Time In"), start)
        if date_matches and start_matches:
            return row
    return None


def validate_xlsx(path: Path) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    workbook = load_workbook(path)
    sheet = workbook["Field Register"]
    headers, rows, table_ref = read_xlsx_rows(path)
    failures: list[dict[str, object]] = []
    if headers != XLSX_HEADERS:
        failures.append({"type": "xlsx_headers", "actual": headers})
    if table_ref != "A4:H8":
        failures.append({"type": "xlsx_table_ref", "actual": table_ref})
    if sheet["A1"].fill.fgColor.rgb != "001F4E78":
        failures.append({"type": "xlsx_title_style", "actual": sheet["A1"].fill.fgColor.rgb})
    if sheet.cell(row=7, column=1).fill.fgColor.rgb != sheet.cell(row=6, column=1).fill.fgColor.rgb:
        failures.append({"type": "xlsx_appended_style_not_copied"})
    for scenario in XLSX_SCENARIOS:
        expected = scenario["expected"]
        actual = find_xlsx_row(rows, expected["Work Date"], expected["Time In"])
        if actual is None:
            failures.append({"type": "xlsx_missing", "scenario": scenario["name"]})
            continue
        for key, value in expected.items():
            if not contains_expected(actual.get(key), value):
                failures.append(
                    {
                        "type": "xlsx_value",
                        "scenario": scenario["name"],
                        "key": key,
                        "expected": value,
                        "actual": actual.get(key),
                    }
                )
    return rows, failures


def validate_docx(path: Path) -> tuple[list[list[str]], list[dict[str, object]]]:
    document = Document(path)
    table = document.tables[0]
    rows = [[cell.text for cell in row.cells] for row in table.rows]
    failures: list[dict[str, object]] = []
    expected_headers = [
        "Asset Tag",
        "Work Date",
        "Person(s)",
        "Action Taken",
        "Result",
        "Remarks",
    ]
    if rows[1] != expected_headers:
        failures.append({"type": "docx_headers", "actual": rows[1]})
    for scenario in DOCX_SCENARIOS:
        expected = scenario["expected"]
        matches = [row for row in rows[2:] if norm(row[0]) == norm(expected[0])]
        if not matches:
            failures.append({"type": "docx_missing", "scenario": scenario["name"]})
            continue
        actual = matches[0]
        for actual_value, expected_value in zip(actual, expected, strict=False):
            if not contains_expected(actual_value, expected_value):
                failures.append(
                    {
                        "type": "docx_value",
                        "scenario": scenario["name"],
                        "expected": expected_value,
                        "actual": actual_value,
                    }
                )
    return rows, failures


class ScenarioNormalizer:
    async def parse_chat_event(self, event, *, media_extractions=None):
        scenario = SCENARIO_BY_SID[event.platform_message_id]
        if scenario in XLSX_SCENARIOS:
            return ChatParseResult(
                intent="document_update",
                document_update_request=DocumentTableUpdateRequest(
                    instruction=str(scenario["body"]),
                    target_document="Styled Field Activity Register Excel",
                    document_kind="xlsx",
                    sheet_name="Daily Activity Log",
                    key_columns=["Date", "Start Time"],
                    rows=[scenario_row_for_request(scenario)],
                    create_if_missing=False,
                ),
                summary_for_user="I will update the styled Excel register.",
                needs_user_confirmation=False,
                confidence=1.0,
            )
        return ChatParseResult(
            intent="document_update",
            document_update_request=DocumentTableUpdateRequest(
                instruction=str(scenario["body"]),
                target_document="Transformer Closeout Word Register",
                document_kind="docx",
                table_name="Transformer Maintenance Table",
                key_columns=["Asset ID"],
                rows=[scenario_row_for_request(scenario)],
                create_if_missing=False,
            ),
            summary_for_user="I will update the Transformer Closeout Word table.",
            needs_user_confirmation=False,
            confidence=1.0,
        )

    async def aclose(self) -> None:
        return None


def scenario_row_for_request(scenario: dict[str, object]) -> dict[str, object]:
    if scenario in XLSX_SCENARIOS:
        expected = scenario["expected"]
        assert isinstance(expected, dict)
        return {
            "Date": expected["Work Date"],
            "Start Time": expected["Time In"],
            "End Time": expected["Time Out"],
            "Activity": expected["Task / Work Done"],
            "People": expected["Technician(s)"],
            "Site": expected["Work Area"],
            "Status": expected["Completion"],
            "Notes": expected["Remarks"],
        }
    expected = scenario["expected"]
    assert isinstance(expected, list)
    return {
        "Asset ID": expected[0],
        "Date": expected[1],
        "People": expected[2],
        "Activity": expected[3],
        "Status": expected[4],
        "Notes": expected[5],
    }


SCENARIO_BY_SID = {scenario["sid"]: scenario for scenario in [*XLSX_SCENARIOS, *DOCX_SCENARIOS]}


def post_twilio_with_retry(client: TestClient, scenario: dict[str, object]):
    last_exc: Exception | None = None
    for attempt in range(1, 4):
        try:
            return client.post(
                "/webhooks/twilio/whatsapp",
                data={
                    "MessageSid": scenario["sid"],
                    "From": "whatsapp:+2348012345678",
                    "WaId": "2348012345678",
                    "Body": scenario["body"],
                    "NumMedia": "0",
                },
            )
        except RuntimeError as exc:
            last_exc = exc
            if "HTTP 503" not in str(exc) or attempt == 3:
                raise
            time.sleep(2 * attempt)
    raise RuntimeError("Twilio post failed after retries") from last_exc


def main() -> None:
    with TemporaryDirectory(prefix="doceebot-complex-pipeline-") as tmp:
        tmp_path = Path(tmp)
        storage_dir = tmp_path / "storage"
        xlsx_seed = tmp_path / "styled_field_register_seed.xlsx"
        docx_seed = tmp_path / "transformer_closeout_seed.docx"
        create_complex_workbook(xlsx_seed)
        create_complex_docx(docx_seed)

        engine = create_engine(f"sqlite+pysqlite:///{tmp_path / 'pipeline.sqlite'}")
        Base.metadata.create_all(engine)
        Session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
        db_session = Session()

        org = Organization(name="Complex Pipeline Test Org")
        user = User(display_name="Samuel Test Worker", phone_number="+2348012345678")
        db_session.add_all([org, user])
        db_session.flush()
        db_session.add(Membership(org_id=org.id, user_id=user.id, role="worker"))
        db_session.commit()

        settings = Settings(
            app_env="development",
            local_storage_dir=str(storage_dir),
            media_storage_backend="local",
            twilio_webhook_auth_enabled=False,
            _env_file=".env",
        )
        app = create_app(settings)

        def override_db_session():
            try:
                yield db_session
            finally:
                pass

        app.dependency_overrides[get_db_session] = override_db_session
        app.dependency_overrides[get_settings] = lambda: settings
        client = TestClient(app)

        with xlsx_seed.open("rb") as handle:
            xlsx_upload = client.post(
                "/dashboard/documents/upload",
                data={
                    "org_id": str(org.id),
                    "owner_user_id": str(user.id),
                    "display_name": "Styled Field Activity Register Excel",
                    "summary": (
                        "Styled Excel field activity register. Real header row is row 4. "
                        "Columns are Work Date, Time In, Time Out, Task / Work Done, "
                        "Technician(s), Work Area, Completion, Remarks."
                    ),
                    "tags": "styled,field-register,excel,complex",
                },
                files={
                    "file": (
                        "styled_field_register.xlsx",
                        handle,
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
                },
            )
        xlsx_upload.raise_for_status()
        xlsx_document = xlsx_upload.json()

        with docx_seed.open("rb") as handle:
            docx_upload = client.post(
                "/dashboard/documents/upload",
                data={
                    "org_id": str(org.id),
                    "owner_user_id": str(user.id),
                    "display_name": "Transformer Closeout Word Register",
                    "summary": (
                        "Word document containing Transformer Maintenance Table. "
                        "The real table headers are on the second row."
                    ),
                    "tags": "transformer,word,closeout,complex",
                },
                files={
                    "file": (
                        "transformer_closeout.docx",
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        docx_upload.raise_for_status()
        docx_document = docx_upload.json()

        original_deepseek_client = chat_processing.DeepSeekClient
        chat_processing.DeepSeekClient = lambda: ScenarioNormalizer()
        scenario_results = []
        try:
            for scenario in [*XLSX_SCENARIOS, *DOCX_SCENARIOS]:
                response = post_twilio_with_retry(client, scenario)
                scenario_results.append(
                    {
                        "name": scenario["name"],
                        "status_code": response.status_code,
                        "twiml": response.text,
                    }
                )
                response.raise_for_status()
        finally:
            chat_processing.DeepSeekClient = original_deepseek_client

        xlsx_final = storage_dir / xlsx_document["storage_key"]
        docx_final = storage_dir / docx_document["storage_key"]
        xlsx_rows, xlsx_failures = validate_xlsx(xlsx_final)
        docx_rows, docx_failures = validate_docx(docx_final)
        failures = [*xlsx_failures, *docx_failures]

        artifact_dir = Path("/root/Doceebot/tmp_pipeline_artifacts")
        artifact_dir.mkdir(parents=True, exist_ok=True)
        xlsx_seed_artifact = artifact_dir / "complex_field_register_before_api_pipeline.xlsx"
        xlsx_final_artifact = artifact_dir / "complex_field_register_after_api_pipeline.xlsx"
        docx_seed_artifact = artifact_dir / "transformer_closeout_before_api_pipeline.docx"
        docx_final_artifact = artifact_dir / "transformer_closeout_after_api_pipeline.docx"
        shutil.copyfile(xlsx_seed, xlsx_seed_artifact)
        shutil.copyfile(xlsx_final, xlsx_final_artifact)
        shutil.copyfile(docx_seed, docx_seed_artifact)
        shutil.copyfile(docx_final, docx_final_artifact)

        report = {
            "xlsx_document_id": xlsx_document["id"],
            "docx_document_id": docx_document["id"],
            "artifacts": {
                "xlsx_before": str(xlsx_seed_artifact),
                "xlsx_after": str(xlsx_final_artifact),
                "docx_before": str(docx_seed_artifact),
                "docx_after": str(docx_final_artifact),
            },
            "scenario_responses": scenario_results,
            "xlsx_rows": xlsx_rows,
            "docx_rows": docx_rows,
            "failures": failures,
            "passed": not failures,
        }
        print(json.dumps(report, indent=2, default=str))
        if failures:
            raise SystemExit(1)


if __name__ == "__main__":
    main()
