from datetime import UTC, datetime
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from whatsapp_ai_agent.config import Settings
from whatsapp_ai_agent.core.events import InboundEvent
from whatsapp_ai_agent.db.models import Base, ManagedDocument
from whatsapp_ai_agent.documents.automation import (
    apply_table_update_to_document,
    automation_catalog,
    create_managed_table_document,
    store_uploaded_document,
)
from whatsapp_ai_agent.documents.schemas import DocumentTableUpdateRequest
from whatsapp_ai_agent.llm.schemas import ChatParseResult
from whatsapp_ai_agent.workflows.chat_processing import process_inbound_event


@pytest.fixture
def db_session():
    engine = create_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    with Session() as session:
        yield session


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Machines"
    sheet.append(["Machine", "Status", "Last Service", "Notes"])
    sheet.append(["Machine X", "Running", "2026-07-01", "Initial commissioning"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Machine Maintenance", level=1)
    table = document.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Machine"
    table.rows[0].cells[1].text = "Status"
    table.rows[0].cells[2].text = "Last Service"
    table.rows[0].cells[3].text = "Notes"
    row = table.add_row().cells
    row[0].text = "Machine X"
    row[1].text = "Running"
    row[2].text = "2026-07-01"
    row[3].text = "Initial commissioning"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _daily_activity_log_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Daily Log"
    sheet.append(
        [
            "Date",
            "Start Time",
            "End Time",
            "Activity",
            "People Participated",
            "Site",
            "Status",
            "Notes",
        ]
    )
    sheet.append(
        [
            "2026-06-10",
            "09:00",
            "11:30",
            "Initial site inspection",
            "Samuel, Ada",
            "Main Plant",
            "done",
            "Baseline inspection completed",
        ]
    )
    sheet.append(
        [
            "2026-06-14",
            "10:00",
            "15:00",
            "Cable tray installation",
            "Samuel, Tunde",
            "Warehouse",
            "in_progress",
            "Covers pending",
        ]
    )
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _complex_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Field Register"
    workbook.create_sheet("Instructions")
    sheet.merge_cells("A1:H1")
    sheet["A1"] = "ABUAD Engineering Field Activity Register"
    sheet["A1"].font = Font(bold=True, color="FFFFFF", size=14)
    sheet["A1"].fill = PatternFill("solid", fgColor="1F4E78")
    sheet["A1"].alignment = Alignment(horizontal="center")
    sheet["A2"] = "The real editable table starts on row 4."
    headers = [
        "Work Date",
        "Time In",
        "Time Out",
        "Task / Work Done",
        "Technician(s)",
        "Work Area",
        "Completion",
        "Remarks",
    ]
    sheet.append([])
    sheet.append(headers)
    rows = [
        [
            "2026-06-20",
            "08:00",
            "10:00",
            "Panel inspection",
            "Samuel, Ada",
            "LV Room",
            "done",
            "No abnormal heating",
        ],
        [
            "2026-06-25",
            "13:00",
            "15:30",
            "Cable ladder alignment",
            "Tunde",
            "Workshop",
            "in_progress",
            "Brackets pending",
        ],
    ]
    thin = Side(style="thin", color="A6A6A6")
    for row in rows:
        sheet.append(row)
    for cell in sheet[4]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="305496")
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(bottom=thin)
    for row in sheet.iter_rows(min_row=5, max_row=6, min_col=1, max_col=8):
        for cell in row:
            cell.fill = PatternFill("solid", fgColor="D9EAF7")
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    table = Table(displayName="FieldRegisterTable", ref="A4:H6")
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    sheet.add_table(table)
    sheet.freeze_panes = "A5"
    sheet.auto_filter.ref = "A4:H6"
    for column, width in zip("ABCDEFGH", [14, 12, 12, 28, 22, 18, 14, 30], strict=False):
        sheet.column_dimensions[column].width = width
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _complex_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Maintenance Closeout Register", level=1)
    document.add_paragraph("The table below has a title row before the actual headers.")
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    title_cells = table.rows[0].cells
    title_cells[0].text = "Transformer Maintenance Table"
    for cell in title_cells[1:]:
        title_cells[0].merge(cell)
    title_cells[0].paragraphs[0].runs[0].font.bold = True
    title_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    header_cells = table.add_row().cells
    for index, header in enumerate(
        ["Asset Tag", "Work Date", "Person(s)", "Action Taken", "Result", "Remarks"]
    ):
        header_cells[index].text = header
        header_cells[index].paragraphs[0].runs[0].font.bold = True
    first = table.add_row().cells
    for index, value in enumerate(
        ["TX-01", "2026-06-19", "Samuel", "Oil level checked", "Open", "Top-up pending"]
    ):
        first[index].text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_uploaded_xlsx_can_be_found_and_updated_by_machine_key(tmp_path, db_session):
    org_id = uuid4()
    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)
    registration = store_uploaded_document(
        org_id=org_id,
        filename="machine-x-register.xlsx",
        data=_xlsx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        summary="Machine X maintenance register",
        tags=["machine-x", "maintenance"],
        db_session=db_session,
        settings=settings,
    )

    result = apply_table_update_to_document(
        document=registration.document,
        request=DocumentTableUpdateRequest(
            instruction="Update Machine X after today's service",
            target_document="machine x register",
            document_kind="xlsx",
            sheet_name="Machines",
            key_columns=["Machine"],
            rows=[
                {
                    "Machine": "Machine X",
                    "Status": "Serviced",
                    "Last Service": "2026-07-02",
                    "Notes": "Oil changed and vibration checked",
                }
            ],
        ),
        db_session=db_session,
        settings=settings,
    )

    assert result.action == "updated"
    assert result.rows_applied == 1
    workbook = load_workbook(registration.document.local_path)
    rows = list(workbook["Machines"].iter_rows(values_only=True))
    assert len(rows) == 2
    assert rows[1] == (
        "Machine X",
        "Serviced",
        "2026-07-02",
        "Oil changed and vibration checked",
    )


def test_uploaded_docx_table_can_be_updated_without_destroying_table(tmp_path, db_session):
    org_id = uuid4()
    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)
    registration = store_uploaded_document(
        org_id=org_id,
        filename="machine-x-log.docx",
        data=_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        summary="Word table for machine maintenance",
        db_session=db_session,
        settings=settings,
    )

    result = apply_table_update_to_document(
        document=registration.document,
        request=DocumentTableUpdateRequest(
            instruction="Update Machine X Word table",
            target_document="machine x log",
            document_kind="docx",
            key_columns=["Machine"],
            rows=[
                {
                    "Machine": "Machine X",
                    "Status": "Needs review",
                    "Last Service": "2026-07-02",
                    "Notes": "Temperature was higher than normal",
                }
            ],
        ),
        db_session=db_session,
        settings=settings,
    )

    assert result.action == "updated"
    document = Document(registration.document.local_path)
    table = document.tables[0]
    assert len(table.rows) == 2
    assert [cell.text for cell in table.rows[1].cells] == [
        "Machine X",
        "Needs review",
        "2026-07-02",
        "Temperature was higher than normal",
    ]


def test_xlsx_update_uses_sheet_fallback_and_header_aliases(tmp_path, db_session):
    org_id = uuid4()
    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)
    registration = store_uploaded_document(
        org_id=org_id,
        filename="daily-activity-log.xlsx",
        data=_daily_activity_log_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        summary="Daily activity log",
        db_session=db_session,
        settings=settings,
    )

    result = apply_table_update_to_document(
        document=registration.document,
        request=DocumentTableUpdateRequest(
            instruction="Update daily activity log row",
            target_document="Daily Activity Log",
            document_kind="xlsx",
            sheet_name="Daily Activity Log",
            key_columns=["Date", "Start Time"],
            rows=[
                {
                    "Date": "2026-06-14",
                    "Start Time": "10:00",
                    "End Time": "15:00",
                    "Activity": "cable tray installation",
                    "People": "Samuel, Tunde and Kemi",
                    "Location": "Warehouse",
                    "Status": "done",
                    "Notes": "added missing trunking covers",
                },
                {
                    "Date": "2026-07-02",
                    "Start Time": "08:15",
                    "End Time": "10:45",
                    "Activity": "generator room inspection",
                    "Participants": "Samuel and Musa",
                    "Location": "Main Plant",
                    "Status": "done",
                    "Notes": "no fault found",
                },
            ],
        ),
        db_session=db_session,
        settings=settings,
    )

    assert result.rows_applied == 2
    workbook = load_workbook(registration.document.local_path)
    sheet = workbook["Daily Log"]
    headers = [cell.value for cell in sheet[1]]
    assert headers == [
        "Date",
        "Start Time",
        "End Time",
        "Activity",
        "People Participated",
        "Site",
        "Status",
        "Notes",
    ]
    rows = list(sheet.iter_rows(min_row=2, values_only=True))
    assert rows[1] == (
        "2026-06-14",
        "10:00",
        "15:00",
        "cable tray installation",
        "Samuel, Tunde and Kemi",
        "Warehouse",
        "done",
        "added missing trunking covers",
    )
    assert rows[2] == (
        "2026-07-02",
        "08:15",
        "10:45",
        "generator room inspection",
        "Samuel and Musa",
        "Main Plant",
        "done",
        "no fault found",
    )


def test_complex_xlsx_preserves_styles_and_adapts_header_row(tmp_path, db_session):
    org_id = uuid4()
    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)
    registration = store_uploaded_document(
        org_id=org_id,
        filename="field-register.xlsx",
        data=_complex_xlsx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        summary="Styled field activity register with header row below title rows",
        db_session=db_session,
        settings=settings,
    )

    result = apply_table_update_to_document(
        document=registration.document,
        request=DocumentTableUpdateRequest(
            instruction="Update styled field activity register",
            target_document="Field Activity Register",
            document_kind="xlsx",
            sheet_name="Daily Activity Log",
            key_columns=["Date", "Start Time"],
            rows=[
                {
                    "Date": "2026-06-25",
                    "Start Time": "13:00",
                    "End Time": "15:30",
                    "Activity": "Cable ladder alignment",
                    "People": "Tunde, Samuel",
                    "Site": "Workshop",
                    "Status": "done",
                    "Notes": "All brackets fitted",
                },
                {
                    "Date": "2026-07-09",
                    "Start Time": "09:10",
                    "End Time": "12:40",
                    "Activity": "Earthing continuity test",
                    "Participants": "Samuel, Grace",
                    "Location": "Admin Block",
                    "Completion": "done",
                    "Remarks": "Readings within limit",
                },
            ],
        ),
        db_session=db_session,
        settings=settings,
    )

    assert result.rows_applied == 2
    workbook = load_workbook(registration.document.local_path)
    sheet = workbook["Field Register"]
    assert sheet["A1"].fill.fgColor.rgb == "001F4E78"
    assert [cell.value for cell in sheet[4]] == [
        "Work Date",
        "Time In",
        "Time Out",
        "Task / Work Done",
        "Technician(s)",
        "Work Area",
        "Completion",
        "Remarks",
    ]
    assert sheet.tables["FieldRegisterTable"].ref == "A4:H7"
    assert [sheet.cell(row=6, column=column).value for column in range(1, 9)] == [
        "2026-06-25",
        "13:00",
        "15:30",
        "Cable ladder alignment",
        "Tunde, Samuel",
        "Workshop",
        "done",
        "All brackets fitted",
    ]
    assert [sheet.cell(row=7, column=column).value for column in range(1, 9)] == [
        "2026-07-09",
        "09:10",
        "12:40",
        "Earthing continuity test",
        "Samuel, Grace",
        "Admin Block",
        "done",
        "Readings within limit",
    ]
    assert (
        sheet.cell(row=7, column=1).fill.fgColor.rgb
        == sheet.cell(row=6, column=1).fill.fgColor.rgb
    )


def test_complex_docx_updates_second_row_header_table(tmp_path, db_session):
    org_id = uuid4()
    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)
    registration = store_uploaded_document(
        org_id=org_id,
        filename="transformer-closeout.docx",
        data=_complex_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        summary="Word closeout register with a title row before table headers",
        db_session=db_session,
        settings=settings,
    )

    result = apply_table_update_to_document(
        document=registration.document,
        request=DocumentTableUpdateRequest(
            instruction="Update transformer closeout Word table",
            target_document="Transformer Maintenance Table",
            document_kind="docx",
            key_columns=["Asset ID"],
            rows=[
                {
                    "Asset ID": "TX-01",
                    "Date": "2026-06-19",
                    "People": "Samuel, Kemi",
                    "Activity": "Oil level checked and gasket inspected",
                    "Status": "Closed",
                    "Notes": "Top-up completed",
                },
                {
                    "Asset ID": "TX-02",
                    "Date": "2026-07-08",
                    "People": "Grace",
                    "Activity": "Insulation resistance test",
                    "Status": "Open",
                    "Notes": "Awaiting supervisor review",
                },
            ],
        ),
        db_session=db_session,
        settings=settings,
    )

    assert result.rows_applied == 2
    document = Document(registration.document.local_path)
    table = document.tables[0]
    assert [cell.text for cell in table.rows[1].cells] == [
        "Asset Tag",
        "Work Date",
        "Person(s)",
        "Action Taken",
        "Result",
        "Remarks",
    ]
    assert [cell.text for cell in table.rows[2].cells] == [
        "TX-01",
        "2026-06-19",
        "Samuel, Kemi",
        "Oil level checked and gasket inspected",
        "Closed",
        "Top-up completed",
    ]
    assert [cell.text for cell in table.rows[3].cells] == [
        "TX-02",
        "2026-07-08",
        "Grace",
        "Insulation resistance test",
        "Open",
        "Awaiting supervisor review",
    ]


def test_create_managed_table_document_from_chat_rows(tmp_path, db_session):
    org_id = uuid4()
    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)

    result = create_managed_table_document(
        org_id=org_id,
        request=DocumentTableUpdateRequest(
            instruction="Create an equipment maintenance log",
            target_document="Equipment Maintenance Log",
            document_kind="xlsx",
            key_columns=["Equipment ID"],
            rows=[{"Equipment ID": "HX-204", "Work Performed": "Pressure tested"}],
        ),
        db_session=db_session,
        settings=settings,
    )

    assert result.action == "created"
    document = db_session.query(ManagedDocument).one()
    assert document.source_type == "chat_created"
    workbook = load_workbook(document.local_path)
    rows = list(workbook.active.iter_rows(values_only=True))
    assert rows[0] == ("Equipment ID", "Work Performed")
    assert rows[1] == ("HX-204", "Pressure tested")


def test_automation_catalog_covers_common_documentation_types():
    titles = {idea.title for idea in automation_catalog()}
    assert "Equipment maintenance log" in titles
    assert "Work order register and closeout form" in titles
    assert "Field inspection report" in titles
    assert "Asset register" in titles


class FakeDocumentUpdateNormalizer:
    async def parse_chat_event(self, event, *, media_extractions=None):
        return ChatParseResult(
            intent="document_update",
            document_update_request=DocumentTableUpdateRequest(
                instruction="Create machine file for Machine X",
                target_document="Machine X Maintenance Log",
                document_kind="xlsx",
                key_columns=["Machine"],
                rows=[{"Machine": "Machine X", "Status": "Installed", "Notes": "Ready"}],
                create_if_missing=True,
            ),
            summary_for_user="I will update the machine documentation.",
            needs_user_confirmation=False,
            confidence=0.9,
        )


@pytest.mark.asyncio
async def test_chat_processing_document_update_intent_creates_file(tmp_path, db_session):
    org_id = uuid4()
    user_id = uuid4()
    event = InboundEvent(
        platform="telegram",
        platform_message_id="chat-1",
        platform_user_id="user-1",
        platform_chat_id="chat",
        message_type="text",
        text="Create machine file for Machine X. Status installed and ready.",
        received_at=datetime(2026, 7, 2, 10, 0, tzinfo=UTC),
        local_date="2026-07-02",
        local_time="11:00:00",
        timezone="Africa/Lagos",
        raw_payload={},
    )

    settings = Settings(local_storage_dir=str(tmp_path), _env_file=None)
    from whatsapp_ai_agent.documents import automation

    original_get_media_storage = automation.get_media_storage
    automation.get_media_storage = lambda _settings=None: original_get_media_storage(settings)
    try:
        result = await process_inbound_event(
            event,
            org_id=org_id,
            user_id=user_id,
            normalizer=FakeDocumentUpdateNormalizer(),
            db_session=db_session,
            store_reports=False,
        )
    finally:
        automation.get_media_storage = original_get_media_storage

    assert result.document_results
    assert "Document automation" in result.reply_text
    assert db_session.query(ManagedDocument).count() == 1
