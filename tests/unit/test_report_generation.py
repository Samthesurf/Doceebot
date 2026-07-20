from datetime import date

import pytest
from docx import Document
from openpyxl import load_workbook

from whatsapp_ai_agent.config import Settings
from whatsapp_ai_agent.documents.reports import deterministic_report_spec, generate_report_files
from whatsapp_ai_agent.llm.prompts import report_spec_user_prompt
from whatsapp_ai_agent.llm.schemas import ReportRequest, WorkLogDraft


class FakeS3Client:
    def __init__(self) -> None:
        self.objects = {}

    def put_object(self, **kwargs):
        self.objects[(kwargs["Bucket"], kwargs["Key"])] = kwargs
        return {"ETag": '"report-etag"'}

    def generate_presigned_url(self, operation, Params, ExpiresIn):
        return f"https://r2.example/{operation}/{Params['Bucket']}/{Params['Key']}?exp={ExpiresIn}"


def sample_logs():
    return [
        WorkLogDraft(
            work_date=date(2026, 7, 1),
            project="Lekki inverter room",
            site="Lekki branch",
            title="DB dressing and continuity test",
            description="The distribution board was dressed and breaker continuity was tested.",
            actions_taken=["Dressed the DB", "Tested breaker continuity"],
            materials_used=["Cable ties", "Labels"],
            status="done",
            confidence=0.9,
        )
    ]


def test_deterministic_report_spec_uses_work_log_details():
    spec = deterministic_report_spec(sample_logs())

    assert spec.title == "Work Report for 2026-07-01"
    joined = "\n".join(paragraph for section in spec.sections for paragraph in section.paragraphs)
    assert "Dressed the DB" in joined
    assert "Lekki branch" in joined


@pytest.mark.asyncio
async def test_generate_report_files_creates_docx_and_xlsx_without_storage(tmp_path):
    files = await generate_report_files(
        org_id="org-1",
        work_logs=sample_logs(),
        output_dir=tmp_path,
        request=ReportRequest(output_format="both"),
        store=False,
        use_llm=False,
    )

    assert {file.format for file in files} == {"docx", "xlsx"}
    docx_file = next(file for file in files if file.format == "docx")
    xlsx_file = next(file for file in files if file.format == "xlsx")

    document = Document(str(docx_file.path))
    paragraphs = "\n".join(p.text for p in document.paragraphs)
    assert "DB dressing and continuity test" in paragraphs
    assert "Dressed the DB" in paragraphs

    workbook = load_workbook(xlsx_file.path)
    rows = list(workbook.active.iter_rows(values_only=True))
    assert rows[0] == ("Date", "Worker", "Project", "Summary")
    assert rows[1][0] == "2026-07-01"
    assert rows[1][2] == "Lekki inverter room"


def test_report_prompt_requests_humanized_weekly_voice():
    prompt = report_spec_user_prompt(
        sample_logs(),
        ReportRequest(report_type="weekly", output_format="docx"),
    )

    assert "natural, human, workmanlike" in prompt
    assert "generic AI phrasing" in prompt
    assert "em dash" in prompt


@pytest.mark.asyncio
async def test_generated_report_uses_twelve_point_times_new_roman(tmp_path):
    files = await generate_report_files(
        org_id="org-1",
        work_logs=sample_logs(),
        output_dir=tmp_path,
        request=ReportRequest(report_type="weekly", output_format="docx"),
        store=False,
        use_llm=False,
    )
    document = Document(str(files[0].path))
    # Only assert font rules on text runs; the logo is an embedded picture run
    # with no text font/size of its own.
    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
    text_runs = [run for run in runs if run.text.strip()]

    assert runs
    assert all(run.font.name == "Times New Roman" for run in text_runs)
    assert all(run.font.size is not None and run.font.size.pt == 12 for run in text_runs)
    assert all(run.font.color is not None and run.font.color.rgb == (0, 0, 0) for run in text_runs)


@pytest.mark.asyncio
async def test_generate_report_files_can_store_to_r2(monkeypatch, tmp_path):
    fake_s3 = FakeS3Client()

    from whatsapp_ai_agent.documents import reports

    settings = Settings(
        media_storage_backend="r2",
        cloudflare_account_id="account-1",
        cloudflare_r2_bucket="doceebot-storage",
        cloudflare_r2_access_key_id="access",
        cloudflare_r2_secret_access_key="secret",
        _env_file=None,
    )

    def fake_storage(settings_arg):
        from whatsapp_ai_agent.media.storage import R2Storage

        return R2Storage(settings_arg, s3_client=fake_s3)

    monkeypatch.setattr(reports, "get_media_storage", fake_storage)

    files = await generate_report_files(
        org_id="org-1",
        work_logs=sample_logs(),
        output_dir=tmp_path,
        request=ReportRequest(output_format="docx"),
        store=True,
        use_llm=False,
        settings=settings,
    )

    assert files[0].stored is not None
    assert files[0].stored.backend == "r2"
    assert files[0].stored.key.startswith("orgs/org-1/generated/")
    assert fake_s3.objects
