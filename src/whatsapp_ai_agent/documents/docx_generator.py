from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from whatsapp_ai_agent.documents.schemas import ReportSpec

_FONT_NAME = "Times New Roman"
_FONT_SIZE = Pt(12)
_BLACK = RGBColor(0, 0, 0)

_LOGO_PATH = Path(__file__).parent / "assets" / "doceebot_logo.png"
_LOGO_WIDTH = Inches(1.0)


def _force_font(style) -> None:  # noqa: ANN001 - python-docx style type is dynamic
    style.font.name = _FONT_NAME
    style.font.size = _FONT_SIZE
    style.font.color.rgb = _BLACK
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attribute}"), _FONT_NAME)
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    for attribute in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attribute}"), None)
    color.set(qn("w:val"), "000000")


def _force_run_font(run) -> None:  # noqa: ANN001 - python-docx run type is dynamic
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE
    run.font.color.rgb = _BLACK
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attribute}"), _FONT_NAME)
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    for attribute in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attribute}"), None)
    color.set(qn("w:val"), "000000")


def _style_document(document) -> None:  # noqa: ANN001 - python-docx object is dynamic
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        _force_font(document.styles[style_name])

    normal = document.styles["Normal"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def _style_paragraph(paragraph, *, heading: bool = False) -> None:  # noqa: ANN001
    paragraph.paragraph_format.space_after = Pt(6 if not heading else 4)
    paragraph.paragraph_format.line_spacing = 1.15
    if heading:
        paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        _force_run_font(run)
        if heading:
            run.font.bold = True


def _clean_text(value: str) -> str:
    return value.replace("—", "-").replace("–", "-")


def generate_docx_report(spec: ReportSpec, output_path: Path) -> Path:
    document = Document()
    _style_document(document)
    if _LOGO_PATH.exists():
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(_LOGO_PATH), width=_LOGO_WIDTH)

    title_text = _clean_text(spec.title)
    document.core_properties.title = title_text
    document.core_properties.author = "Doceebot"
    document.core_properties.subject = "Weekly work report"

    title = document.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_paragraph(title, heading=True)

    for section in spec.sections:
        heading = document.add_heading(_clean_text(section.heading), level=2)
        _style_paragraph(heading, heading=True)
        for text in section.paragraphs:
            paragraph = document.add_paragraph(_clean_text(text))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _style_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
