import json
import re
import shutil
from copy import copy, deepcopy
from dataclasses import dataclass
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any
from uuid import UUID, uuid4

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter, range_boundaries
from sqlalchemy.orm import Session

from whatsapp_ai_agent.config import Settings, get_settings
from whatsapp_ai_agent.core.events import InboundEvent, MediaRef
from whatsapp_ai_agent.db.models import ManagedDocument, RawInboundMessage
from whatsapp_ai_agent.db.repositories import ManagedDocumentRepository
from whatsapp_ai_agent.documents.schemas import (
    DocumentationAutomationIdea,
    DocumentAutomationResult,
    DocumentKind,
    DocumentTableUpdateRequest,
    ManagedDocumentSummary,
)
from whatsapp_ai_agent.media.storage import StoredObject, get_media_storage, org_object_key

SPREADSHEET_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_KIND_SUFFIX = {
    "xlsx": ".xlsx",
    "docx": ".docx",
    "csv": ".csv",
    "pdf": ".pdf",
    "text": ".txt",
}

_KIND_CONTENT_TYPE = {
    "xlsx": SPREADSHEET_CONTENT_TYPE,
    "docx": DOCX_CONTENT_TYPE,
    "csv": "text/csv",
    "pdf": "application/pdf",
    "text": "text/plain; charset=utf-8",
}

DOCUMENTATION_AUTOMATION_CATALOG: tuple[DocumentationAutomationIdea, ...] = (
    DocumentationAutomationIdea(
        slug="equipment-maintenance-log",
        title="Equipment maintenance log",
        formats=["xlsx", "docx"],
        chat_inputs=[
            "Machine ID or asset name",
            "date and technician",
            "work done",
            "parts replaced",
            "cost or downtime",
            "next service due",
        ],
        typical_fields=[
            "Equipment ID",
            "Location",
            "Service date",
            "Technician",
            "Work performed",
            "Parts used",
            "Cost",
            "Next service due",
        ],
        automation_notes=(
            "Best handled as a longitudinal Excel workbook or Word table. Chat updates can "
            "upsert by Equipment ID and append service-history rows."
        ),
        source_notes=(
            "Maintenance documentation references commonly separate asset master records, "
            "work orders, inspection reports, service records, and failure analysis."
        ),
    ),
    DocumentationAutomationIdea(
        slug="work-order-register",
        title="Work order register and closeout form",
        formats=["xlsx", "docx"],
        chat_inputs=[
            "work order number",
            "request description",
            "assigned technician",
            "priority",
            "materials and labor hours",
            "completion notes",
        ],
        typical_fields=[
            "Work Order No",
            "Asset",
            "Priority",
            "Assigned To",
            "Date Started",
            "Date Completed",
            "Materials",
            "Labor Hours",
            "Closeout Status",
        ],
        automation_notes=(
            "The bot can create one row per work order, then update status, labor, parts, "
            "verification tests, and closeout notes as messages arrive."
        ),
        source_notes=(
            "Field service and maintenance templates emphasize request, assignment, "
            "materials, labor, completion, and authorization sections."
        ),
    ),
    DocumentationAutomationIdea(
        slug="field-inspection-report",
        title="Field inspection report",
        formats=["xlsx", "docx"],
        chat_inputs=[
            "site or asset",
            "inspection type",
            "measurements with units",
            "condition found",
            "photos or references",
            "recommended action and priority",
        ],
        typical_fields=[
            "Inspection Date",
            "Inspector",
            "Asset/Site",
            "Inspection Type",
            "Measured Values",
            "Condition",
            "Photos",
            "Recommendation",
            "Priority",
        ],
        automation_notes=(
            "Works well from voice notes because the system can extract measured values, "
            "risk level, and recommendations into a repeatable table."
        ),
        source_notes=(
            "Inspection report templates usually capture job details, observations, "
            "measurements, condition checks, recommendations, and photo references."
        ),
    ),
    DocumentationAutomationIdea(
        slug="asset-register",
        title="Asset register",
        formats=["xlsx"],
        chat_inputs=[
            "asset name or tag",
            "serial number",
            "location",
            "purchase or installation date",
            "warranty details",
            "current condition",
        ],
        typical_fields=[
            "Asset ID",
            "Description",
            "Serial No",
            "Location",
            "Owner",
            "Install Date",
            "Warranty Expiry",
            "Condition",
        ],
        automation_notes=(
            "The bot can upsert by Asset ID or serial number and link manuals, SOPs, "
            "inspection forms, and maintenance history."
        ),
        source_notes=(
            "Asset master records form the foundation for maintenance records, work orders, "
            "and inspections."
        ),
    ),
    DocumentationAutomationIdea(
        slug="preventive-maintenance-schedule",
        title="Preventive maintenance schedule",
        formats=["xlsx"],
        chat_inputs=[
            "asset ID",
            "maintenance interval",
            "last service date",
            "next due date",
            "responsible person",
            "completion status",
        ],
        typical_fields=[
            "Asset ID",
            "Task",
            "Interval",
            "Last Done",
            "Next Due",
            "Assigned To",
            "Status",
        ],
        automation_notes=(
            "A chat message like 'pump P-12 serviced today, next one in 30 days' can update "
            "the last-done date, next-due date, and schedule status."
        ),
        source_notes="Preventive schedules are often paired with maintenance logs and work orders.",
    ),
    DocumentationAutomationIdea(
        slug="sop-or-checklist",
        title="SOP or inspection checklist",
        formats=["docx", "xlsx"],
        chat_inputs=[
            "procedure name",
            "asset or area",
            "steps performed",
            "pass/fail checks",
            "deviations",
            "sign-off details",
        ],
        typical_fields=[
            "Step No",
            "Procedure Step",
            "Expected Result",
            "Actual Result",
            "Pass/Fail",
            "Notes",
            "Verified By",
        ],
        automation_notes=(
            "The system can convert SOP text into a checklist and later fill the checklist "
            "from voice notes while preserving an audit trail."
        ),
        source_notes=(
            "Modern SOP guidance recommends embedding SOP steps into digital forms and "
            "recording completion evidence."
        ),
    ),
    DocumentationAutomationIdea(
        slug="non-conformance-corrective-action",
        title="Non-conformance and corrective action register",
        formats=["xlsx", "docx"],
        chat_inputs=[
            "non-conformance description",
            "location",
            "root cause",
            "corrective action",
            "owner",
            "due date",
            "closure evidence",
        ],
        typical_fields=[
            "NCR No",
            "Date",
            "Location",
            "Issue",
            "Root Cause",
            "Corrective Action",
            "Owner",
            "Due Date",
            "Status",
        ],
        automation_notes=(
            "Useful for facility maintenance, quality, and safety teams. Each chat update "
            "can move an item from open to closed with evidence."
        ),
        source_notes=(
            "Facility maintenance documentation often includes non-conformance reports, "
            "corrective action requests, preventive action requests, and defect notices."
        ),
    ),
)


@dataclass(frozen=True)
class StoredDocumentRegistration:
    document: ManagedDocument
    stored: StoredObject


def infer_document_kind(filename: str | None, content_type: str | None = None) -> DocumentKind:
    suffix = Path(filename or "").suffix.lower()
    content = (content_type or "").lower()
    if suffix in {".xlsx", ".xlsm"} or "spreadsheetml" in content or "excel" in content:
        return "xlsx"
    if suffix == ".docx" or "wordprocessingml" in content or "msword" in content:
        return "docx"
    if suffix == ".csv" or "text/csv" in content:
        return "csv"
    if suffix == ".pdf" or "application/pdf" in content:
        return "pdf"
    if suffix in {".txt", ".md"} or content.startswith("text/"):
        return "text"
    return "unknown"


def content_type_for_kind(kind: str, fallback: str | None = None) -> str | None:
    return fallback or _KIND_CONTENT_TYPE.get(kind)


def safe_filename(filename: str, *, fallback: str = "document") -> str:
    name = Path(filename).name.strip() or fallback
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name or fallback


def document_summary_from_model(document: ManagedDocument) -> ManagedDocumentSummary:
    try:
        tags = json.loads(document.tags_json or "[]")
    except json.JSONDecodeError:
        tags = []
    if not isinstance(tags, list):
        tags = []
    document_kind = document.document_kind if document.document_kind in _KIND_SUFFIX else "unknown"
    return ManagedDocumentSummary(
        id=str(document.id),
        org_id=str(document.org_id),
        filename=document.filename,
        display_name=document.display_name,
        document_kind=document_kind,
        source_type=document.source_type,
        status=document.status,
        content_type=document.content_type,
        size_bytes=document.size_bytes,
        sha256_hex=document.sha256_hex,
        storage_backend=document.storage_backend,
        storage_key=document.storage_key,
        url=document.storage_url,
        summary=document.summary,
        tags=[str(tag) for tag in tags],
        created_at=document.created_at.isoformat() if document.created_at else None,
        updated_at=document.updated_at.isoformat() if document.updated_at else None,
    )


def store_uploaded_document(
    *,
    org_id: UUID,
    filename: str,
    data: bytes,
    db_session: Session,
    content_type: str | None = None,
    owner_user_id: UUID | None = None,
    display_name: str | None = None,
    summary: str | None = None,
    tags: list[str] | None = None,
    settings: Settings | None = None,
) -> StoredDocumentRegistration:
    settings = settings or get_settings()
    kind = infer_document_kind(filename, content_type)
    clean_name = safe_filename(filename, fallback=f"document{_KIND_SUFFIX.get(kind, '')}")
    key = org_object_key(str(org_id), "managed", kind, f"{uuid4()}-{clean_name}")
    storage = get_media_storage(settings)
    stored = storage.save_bytes(
        key,
        data,
        content_type=content_type_for_kind(kind, content_type),
        metadata={"org_id": str(org_id), "source_type": "managed_document"},
    )
    document = ManagedDocumentRepository(db_session).add_from_stored_object(
        org_id=org_id,
        owner_user_id=owner_user_id,
        stored=stored,
        filename=clean_name,
        document_kind=kind,
        content_type=content_type_for_kind(kind, content_type),
        display_name=display_name or Path(clean_name).stem,
        source_type="uploaded",
        summary=summary,
        tags=tags,
    )
    return StoredDocumentRegistration(document=document, stored=stored)


def register_generated_document_file(
    *,
    org_id: UUID,
    path: Path,
    db_session: Session,
    stored: StoredObject | None = None,
    owner_user_id: UUID | None = None,
    display_name: str | None = None,
    summary: str | None = None,
    settings: Settings | None = None,
) -> ManagedDocument:
    settings = settings or get_settings()
    filename = safe_filename(path.name)
    kind = infer_document_kind(filename, None)
    if stored is None:
        key = org_object_key(str(org_id), "generated", filename)
        storage = get_media_storage(settings)
        stored = storage.save_file(
            key,
            path,
            content_type=content_type_for_kind(kind),
            metadata={"org_id": str(org_id), "source_type": "generated_report"},
        )
    return ManagedDocumentRepository(db_session).add_from_stored_object(
        org_id=org_id,
        owner_user_id=owner_user_id,
        stored=stored,
        filename=filename,
        document_kind=kind,
        content_type=content_type_for_kind(kind, stored.content_type),
        display_name=display_name or Path(filename).stem,
        source_type="generated",
        summary=summary,
    )


def register_pending_inbound_media_document(
    *,
    event: InboundEvent,
    media: MediaRef,
    db_session: Session,
    org_id: UUID,
    owner_user_id: UUID | None = None,
) -> ManagedDocument | None:
    kind = infer_document_kind(media.filename, media.content_type)
    if kind not in {"xlsx", "docx", "csv", "pdf", "text"}:
        return None
    filename = safe_filename(media.filename or f"upload-{media.index}{_KIND_SUFFIX.get(kind, '')}")
    if media.storage_key:
        stored = StoredObject(
            backend=media.storage_backend or event.platform,
            key=media.storage_key,
            content_type=content_type_for_kind(kind, media.content_type),
            size_bytes=media.size_bytes,
            sha256_hex=media.sha256_hex,
            url=media.storage_url,
        )
        status = "available"
        summary = "Uploaded through chat and stored from platform media bytes."
    else:
        key = org_object_key(str(org_id), "pending-uploads", event.platform_message_id, filename)
        stored = StoredObject(
            backend=event.platform,
            key=key,
            content_type=content_type_for_kind(kind, media.content_type),
            size_bytes=media.size_bytes,
            url=media.url,
        )
        status = "pending_download"
        summary = "Uploaded through chat and waiting for media download storage."
    return ManagedDocumentRepository(db_session).add_from_stored_object(
        org_id=org_id,
        owner_user_id=owner_user_id,
        stored=stored,
        filename=filename,
        document_kind=kind,
        content_type=content_type_for_kind(kind, media.content_type),
        display_name=Path(filename).stem,
        source_type="uploaded",
        status=status,
        summary=summary,
        tags=[event.platform, event.message_type],
    )


def create_managed_table_document(
    *,
    org_id: UUID,
    request: DocumentTableUpdateRequest,
    db_session: Session,
    owner_user_id: UUID | None = None,
    settings: Settings | None = None,
) -> DocumentAutomationResult:
    if not request.rows and not request.key_columns:
        raise ValueError("at least one row or key column is required to create a document")
    settings = settings or get_settings()
    kind = request.document_kind if request.document_kind in {"xlsx", "docx"} else "xlsx"
    base_name = request.target_document or request.table_name or "documentation-log"
    filename = safe_filename(base_name, fallback="documentation-log")
    suffix = _KIND_SUFFIX[kind]
    if not filename.lower().endswith(suffix):
        filename = f"{filename}{suffix}"

    with TemporaryDirectory() as tmp_dir:
        output_path = Path(tmp_dir) / filename
        if kind == "xlsx":
            _create_xlsx_table(output_path, request)
        else:
            _create_docx_table(output_path, request)
        key = org_object_key(str(org_id), "managed", kind, f"{uuid4()}-{filename}")
        storage = get_media_storage(settings)
        stored = storage.save_file(
            key,
            output_path,
            content_type=content_type_for_kind(kind),
            metadata={"org_id": str(org_id), "source_type": "managed_document"},
        )

    document = ManagedDocumentRepository(db_session).add_from_stored_object(
        org_id=org_id,
        owner_user_id=owner_user_id,
        stored=stored,
        filename=filename,
        document_kind=kind,
        content_type=content_type_for_kind(kind),
        display_name=Path(filename).stem,
        source_type="chat_created",
        summary=request.instruction or "Created from chat input.",
        tags=["chat-created", "table"],
    )
    changes = [f"Created {kind.upper()} document with {len(request.rows)} row(s)."]
    return DocumentAutomationResult(
        document=document_summary_from_model(document),
        action="created",
        changes=changes,
        rows_applied=len(request.rows),
    )


def apply_table_update_to_document(
    *,
    document: ManagedDocument,
    request: DocumentTableUpdateRequest,
    db_session: Session,
    user_id: UUID | None = None,
    raw_message: RawInboundMessage | None = None,
    settings: Settings | None = None,
) -> DocumentAutomationResult:
    if document.status != "available":
        raise ValueError(f"document is not available for editing yet: {document.status}")
    if document.document_kind not in {"xlsx", "docx"}:
        raise ValueError("only XLSX and DOCX table documents can be updated automatically")
    if not request.rows:
        raise ValueError("at least one row is required for a table update")

    settings = settings or get_settings()
    storage = get_media_storage(settings)
    suffix = _KIND_SUFFIX.get(document.document_kind, Path(document.filename).suffix or ".tmp")
    with TemporaryDirectory() as tmp_dir:
        work_path = Path(tmp_dir) / safe_filename(document.filename, fallback=f"document{suffix}")
        if document.local_path and Path(document.local_path).exists():
            shutil.copyfile(document.local_path, work_path)
        elif hasattr(storage, "read_bytes"):
            work_path.write_bytes(storage.read_bytes(document.storage_key))  # type: ignore[union-attr]
        else:
            raise RuntimeError("configured storage cannot read existing document bytes")

        if document.document_kind == "xlsx":
            changes = append_or_update_xlsx_table(work_path, request)
        else:
            changes = append_or_update_docx_table(work_path, request)

        stored = storage.save_file(
            document.storage_key,
            work_path,
            content_type=content_type_for_kind(document.document_kind, document.content_type),
            metadata={"org_id": str(document.org_id), "source_type": "managed_document"},
        )

    repo = ManagedDocumentRepository(db_session)
    updated = repo.update_storage_info(document, stored)
    repo.add_update_record(
        org_id=document.org_id,
        document_id=document.id,
        user_id=user_id,
        raw_message_id=raw_message.id if raw_message else None,
        instruction=request.instruction,
        changes=changes,
    )
    return DocumentAutomationResult(
        document=document_summary_from_model(updated),
        action="updated",
        changes=changes,
        rows_applied=len(request.rows),
    )


def automation_catalog(industry: str | None = None) -> list[DocumentationAutomationIdea]:
    if not industry:
        return list(DOCUMENTATION_AUTOMATION_CATALOG)
    terms = {term for term in industry.lower().split() if term}
    result: list[DocumentationAutomationIdea] = []
    for idea in DOCUMENTATION_AUTOMATION_CATALOG:
        haystack = " ".join(
            [
                idea.title,
                idea.slug,
                " ".join(idea.typical_fields),
                idea.automation_notes,
                idea.source_notes or "",
            ]
        ).lower()
        if any(term in haystack for term in terms):
            result.append(idea)
    return result or list(DOCUMENTATION_AUTOMATION_CATALOG)


def append_or_update_xlsx_table(path: Path, request: DocumentTableUpdateRequest) -> list[str]:
    workbook = load_workbook(path)
    sheet = _select_xlsx_sheet(workbook, request.sheet_name)
    header_row = _find_xlsx_header_row(sheet, request)
    request = _canonicalize_table_request(request, _current_xlsx_headers(sheet, header_row))
    headers = _ensure_xlsx_headers(sheet, request, header_row)
    header_index = {_normalize_header(header): index + 1 for index, header in enumerate(headers)}
    xlsx_table = _matching_xlsx_table(sheet, header_row)
    changes: list[str] = []
    for row in request.rows:
        row_number = _find_matching_xlsx_row(
            sheet,
            header_index,
            row,
            request.key_columns,
            header_row=header_row,
        )
        if row_number is None:
            row_number = _next_xlsx_insert_row(sheet, header_row, xlsx_table)
            _prepare_appended_xlsx_row(sheet, row_number, header_row, len(headers))
            if xlsx_table is not None:
                _expand_xlsx_table(xlsx_table, row_number, len(headers))
            changes.append(f"Appended row {row_number} to sheet {sheet.title}.")
        else:
            changes.append(f"Updated row {row_number} in sheet {sheet.title}.")
        for key, value in row.items():
            normalized = _normalize_header(key)
            if normalized not in header_index:
                headers.append(str(key))
                column = len(headers)
                sheet.cell(row=header_row, column=column).value = str(key)
                header_index[normalized] = column
            sheet.cell(row=row_number, column=header_index[normalized]).value = _cell_value(value)
    workbook.save(path)
    return changes


def append_or_update_docx_table(path: Path, request: DocumentTableUpdateRequest) -> list[str]:
    document = Document(str(path))
    table, headers, header_row_index = _find_or_create_docx_table(document, request)
    request = _canonicalize_table_request(request, headers)
    changes: list[str] = []
    for row in request.rows:
        headers = _ensure_docx_headers(table, headers, row, header_row_index)
        header_index = {_normalize_header(header): index for index, header in enumerate(headers)}
        row_obj, row_number = _find_matching_docx_row(
            table,
            header_index,
            row,
            request.key_columns,
            header_row_index=header_row_index,
        )
        if row_obj is None:
            row_obj = _append_docx_row_like_existing(table, header_row_index)
            row_number = len(table.rows)
            changes.append(f"Appended row {row_number} to Word table.")
        else:
            changes.append(f"Updated row {row_number} in Word table.")
        for key, value in row.items():
            column = header_index[_normalize_header(key)]
            row_obj.cells[column].text = str(_cell_value(value))
    document.save(str(path))
    return changes


def _create_xlsx_table(path: Path, request: DocumentTableUpdateRequest) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = (request.sheet_name or request.table_name or "Documentation")[:31]
    headers = _headers_from_request(request)
    sheet.append(headers)
    for row in request.rows:
        sheet.append([_cell_value(row.get(header)) for header in headers])
    workbook.save(path)


def _create_docx_table(path: Path, request: DocumentTableUpdateRequest) -> None:
    document = Document()
    heading = request.table_name or request.target_document or "Documentation Log"
    document.add_heading(heading, level=1)
    headers = _headers_from_request(request)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in request.rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(_cell_value(row.get(header)))
    document.save(str(path))


def _headers_from_request(request: DocumentTableUpdateRequest) -> list[str]:
    headers: list[str] = []
    for key in request.key_columns:
        if key and _normalize_header(key) not in {_normalize_header(header) for header in headers}:
            headers.append(key)
    for row in request.rows:
        for key in row:
            if _normalize_header(key) not in {_normalize_header(header) for header in headers}:
                headers.append(str(key))
    return headers or ["Item", "Details"]


_TABLE_HEADER_ALIAS_GROUPS: tuple[set[str], ...] = (
    {"date", "day", "work date", "activity date", "service date", "inspection date"},
    {"start", "start time", "time started", "time in", "from", "start_time"},
    {"end", "end time", "finish time", "time out", "to", "end_time"},
    {
        "activity",
        "task",
        "work",
        "work done",
        "task work done",
        "activity performed",
        "work performed",
        "action taken",
    },
    {
        "people",
        "person",
        "persons",
        "personnel",
        "participant",
        "participants",
        "people participated",
        "people involved",
        "technician",
        "technicians",
        "workers",
        "worker",
        "team",
    },
    {"site", "location", "place", "work site", "work area", "area"},
    {"status", "completion", "completion status", "state", "closeout status", "result"},
    {"notes", "note", "remarks", "remark", "comments", "comment"},
    {"machine", "machine id", "machine name", "equipment", "equipment id"},
    {"asset", "asset id", "asset tag", "asset no", "asset number"},
    {"serial no", "serial number", "s n", "sn"},
    {"work order no", "work order number", "wo no", "work order"},
)


def _current_xlsx_headers(sheet: Any, header_row: int) -> list[str]:
    return [
        str(cell.value).strip()
        for cell in sheet[header_row]
        if cell.value not in {None, ""}
    ]


def _alias_group_for_header(value: Any) -> set[str]:
    normalized = _normalize_header(value)
    for group in _TABLE_HEADER_ALIAS_GROUPS:
        if normalized in group:
            return group
    return {normalized}


def _headers_are_equivalent(left: Any, right: Any) -> bool:
    left_norm = _normalize_header(left)
    right_norm = _normalize_header(right)
    if not left_norm or not right_norm:
        return False
    if left_norm == right_norm:
        return True
    left_group = _alias_group_for_header(left_norm)
    right_group = _alias_group_for_header(right_norm)
    if left_group & right_group:
        return True
    return left_norm in right_norm or right_norm in left_norm


def _canonical_header_for_key(key: str, existing_headers: list[str]) -> str:
    normalized_key = _normalize_header(key)
    exact_headers = {_normalize_header(header): header for header in existing_headers}
    if normalized_key in exact_headers:
        return exact_headers[normalized_key]

    for header in existing_headers:
        if _headers_are_equivalent(key, header):
            return header
    return key


def _canonicalize_table_request(
    request: DocumentTableUpdateRequest,
    existing_headers: list[str],
) -> DocumentTableUpdateRequest:
    if not existing_headers:
        return request

    rows: list[dict[str, Any]] = []
    for row in request.rows:
        canonical_row: dict[str, Any] = {}
        for key, value in row.items():
            canonical_key = _canonical_header_for_key(str(key), existing_headers)
            if canonical_key in canonical_row and canonical_row[canonical_key] not in {None, ""}:
                continue
            canonical_row[canonical_key] = value
        rows.append(canonical_row)

    key_columns = [
        _canonical_header_for_key(str(key), existing_headers) for key in request.key_columns
    ]
    return request.model_copy(update={"rows": rows, "key_columns": key_columns})


def _request_header_candidates(request: DocumentTableUpdateRequest) -> list[str]:
    candidates: list[str] = []
    candidates.extend(str(key) for key in request.key_columns)
    for row in request.rows:
        candidates.extend(str(key) for key in row)
    return [candidate for candidate in candidates if _normalize_header(candidate)]


def _header_match_score(headers: list[Any], request: DocumentTableUpdateRequest) -> int:
    candidates = _request_header_candidates(request)
    score = 0
    for header in headers:
        if not _normalize_header(header):
            continue
        for candidate in candidates:
            if _headers_are_equivalent(candidate, header):
                score += 3 if candidate in request.key_columns else 1
                break
    return score


def _xlsx_table_bounds(table: Any) -> tuple[int, int, int, int]:
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return min_col or 1, min_row or 1, max_col or 1, max_row or 1


def _find_xlsx_header_row(sheet: Any, request: DocumentTableUpdateRequest) -> int:
    best_row = 1
    best_score = -1
    for table in _iter_xlsx_tables(sheet):
        _, min_row, _, _ = _xlsx_table_bounds(table)
        headers = _xlsx_headers_at_row(sheet, min_row)
        score = _header_match_score(headers, request)
        if score > best_score:
            best_row = min_row
            best_score = score

    search_limit = min(sheet.max_row, 25)
    for row_number in range(1, search_limit + 1):
        headers = _xlsx_headers_at_row(sheet, row_number)
        score = _header_match_score(headers, request)
        if score > best_score:
            best_row = row_number
            best_score = score
    return best_row


def _xlsx_headers_at_row(sheet: Any, row_number: int) -> list[str]:
    return [
        str(cell.value).strip() if cell.value not in {None, ""} else ""
        for cell in sheet[row_number]
    ]


def _iter_xlsx_tables(sheet: Any) -> list[Any]:
    tables = getattr(sheet, "tables", {})
    if hasattr(tables, "values"):
        return list(tables.values())
    return list(tables)


def _matching_xlsx_table(sheet: Any, header_row: int) -> Any | None:
    for table in _iter_xlsx_tables(sheet):
        _, min_row, _, _ = _xlsx_table_bounds(table)
        if min_row == header_row:
            return table
    return None


def _next_xlsx_insert_row(sheet: Any, header_row: int, table: Any | None) -> int:
    if table is not None:
        _, _, _, max_row = _xlsx_table_bounds(table)
        row_number = max_row + 1
        if row_number <= sheet.max_row:
            sheet.insert_rows(row_number)
        return row_number
    return max(sheet.max_row + 1, header_row + 1)


def _prepare_appended_xlsx_row(
    sheet: Any,
    row_number: int,
    header_row: int,
    max_columns: int,
) -> None:
    template_row = row_number - 1 if row_number > header_row + 1 else header_row
    if sheet.row_dimensions[template_row].height is not None:
        sheet.row_dimensions[row_number].height = sheet.row_dimensions[template_row].height
    for column in range(1, max_columns + 1):
        source = sheet.cell(row=template_row, column=column)
        target = sheet.cell(row=row_number, column=column)
        if source.has_style:
            target._style = copy(source._style)
        if source.number_format:
            target.number_format = source.number_format
        if source.alignment:
            target.alignment = copy(source.alignment)
        if source.protection:
            target.protection = copy(source.protection)


def _expand_xlsx_table(table: Any, row_number: int, max_columns: int) -> None:
    min_col, min_row, max_col, max_row = _xlsx_table_bounds(table)
    max_col = max(max_col, max_columns)
    max_row = max(max_row, row_number)
    table.ref = f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{max_row}"


def _select_xlsx_sheet(workbook: Any, sheet_name: str | None) -> Any:
    if not sheet_name:
        return workbook.active
    if sheet_name in workbook.sheetnames:
        return workbook[sheet_name]

    requested = _normalize_header(sheet_name)
    for existing_name in workbook.sheetnames:
        if _normalize_header(existing_name) == requested:
            return workbook[existing_name]

    requested_terms = {term for term in requested.split() if term}
    best_name = None
    best_overlap = 0
    for existing_name in workbook.sheetnames:
        existing_terms = {term for term in _normalize_header(existing_name).split() if term}
        overlap = len(requested_terms & existing_terms)
        if overlap > best_overlap:
            best_name = existing_name
            best_overlap = overlap
    if best_name and best_overlap:
        return workbook[best_name]
    return workbook.active


def _ensure_xlsx_headers(
    sheet: Any,
    request: DocumentTableUpdateRequest,
    header_row: int,
) -> list[str]:
    headers = [
        str(cell.value).strip()
        for cell in sheet[header_row]
        if cell.value not in {None, ""}
    ]
    if not headers:
        headers = _headers_from_request(request)
        for index, header in enumerate(headers, start=1):
            sheet.cell(row=header_row, column=index).value = header
    normalized = {_normalize_header(header) for header in headers}
    for row in request.rows:
        for key in row:
            if _normalize_header(key) not in normalized:
                headers.append(str(key))
                column = len(headers)
                source = sheet.cell(row=header_row, column=max(1, column - 1))
                target = sheet.cell(row=header_row, column=column)
                target.value = str(key)
                if source.has_style:
                    target._style = copy(source._style)
                normalized.add(_normalize_header(key))
    return headers


def _find_matching_xlsx_row(
    sheet: Any,
    header_index: dict[str, int],
    row: dict[str, Any],
    key_columns: list[str],
    *,
    header_row: int,
) -> int | None:
    usable_keys = [
        key for key in key_columns if _normalize_header(key) in header_index and key in row
    ]
    if not usable_keys:
        return None
    for row_number in range(header_row + 1, sheet.max_row + 1):
        if all(
            _normalize_cell(
                sheet.cell(row=row_number, column=header_index[_normalize_header(key)]).value
            )
            == _normalize_cell(row.get(key))
            for key in usable_keys
        ):
            return row_number
    return None


def _find_or_create_docx_table(
    document: Any,
    request: DocumentTableUpdateRequest,
) -> tuple[Any, list[str], int]:
    best_match: tuple[int, Any, list[str], int] | None = None
    for table in document.tables:
        if not table.rows:
            continue
        for row_index, row in enumerate(table.rows[:10]):
            headers = [cell.text.strip() for cell in row.cells]
            score = _header_match_score(headers, request)
            if score <= 0:
                continue
            if best_match is None or score > best_match[0]:
                best_match = (score, table, headers, row_index)
    if best_match is not None:
        _, table, headers, row_index = best_match
        return table, headers, row_index

    desired_headers = _headers_from_request(request)
    document.add_paragraph()
    heading = request.table_name or request.target_document or "Documentation Update"
    document.add_heading(heading, level=2)
    table = document.add_table(rows=1, cols=len(desired_headers))
    table.style = "Table Grid"
    for index, header in enumerate(desired_headers):
        table.rows[0].cells[index].text = header
    return table, desired_headers, 0


def _ensure_docx_headers(
    table: Any,
    headers: list[str],
    row: dict[str, Any],
    header_row_index: int,
) -> list[str]:
    normalized = {_normalize_header(header) for header in headers}
    for key in row:
        if _normalize_header(key) not in normalized:
            table.add_column(Inches(1.4))
            headers.append(str(key))
            table.rows[header_row_index].cells[len(headers) - 1].text = str(key)
            normalized.add(_normalize_header(key))
    return headers


def _append_docx_row_like_existing(table: Any, header_row_index: int) -> Any:
    template_index = len(table.rows) - 1
    if template_index <= header_row_index:
        template_index = header_row_index
    template = table.rows[template_index]
    table._tbl.append(deepcopy(template._tr))
    new_row = table.rows[-1]
    for cell in new_row.cells:
        cell.text = ""
    return new_row


def _find_matching_docx_row(
    table: Any,
    header_index: dict[str, int],
    row: dict[str, Any],
    key_columns: list[str],
    *,
    header_row_index: int,
) -> tuple[Any | None, int | None]:
    usable_keys = [
        key for key in key_columns if _normalize_header(key) in header_index and key in row
    ]
    if not usable_keys:
        return None, None
    for index, row_obj in enumerate(table.rows[header_row_index + 1 :], start=header_row_index + 2):
        if all(
            _normalize_cell(row_obj.cells[header_index[_normalize_header(key)]].text)
            == _normalize_cell(row.get(key))
            for key in usable_keys
        ):
            return row_obj, index
    return None, None


def _normalize_header(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = text.replace("(s)", "s").replace("&", " and ")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _normalize_cell(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


def _cell_value(value: Any) -> Any:
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    return json.dumps(value, ensure_ascii=False)
